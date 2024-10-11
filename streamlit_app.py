import streamlit as st
import sys
import os
import anthropic
import numpy as np
from typing import List, Dict
import json
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.shared import Pt
import io
import voyageai


# Initialize clients with API keys from environment variables
anthropic_client = anthropic.Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY'))
voyage_client = voyageai.Client(api_key=os.getenv('VOYAGE_API_KEY'))

# Load pre-computed embeddings and text chunks
try:
    embeddings = np.load('./dsmv_embeddings.npy')
    with open('./dsmv_text_chunks.json', 'r') as f:
        text_chunks = json.load(f)
except FileNotFoundError:
    st.error("Required data files (dsmv_embeddings.npy or dsmv_text_chunks.json) not found.")
    embeddings = None
    text_chunks = None

class DiagnosticTreatmentAgent:
    def __init__(self):
        self.opus_model = "claude-3-5-sonnet-20240620"
        self.haiku_model = "claude-3-5-sonnet-20240620"
        self.document = Document()
        self.document.add_heading('Diagnostic and Treatment Report', 0)

    def add_to_document(self, heading, content, level=1):
        self.document.add_heading(heading, level)
        self.document.add_paragraph(content)

    def get_relevant_dsm5_criteria(self, patient_data: str, initial_top_k: int = 10, final_top_k: int = 5) -> List[Dict]:
        if voyage_client is None:
            st.warning("Voyage client is not available. Skipping DSM-5 criteria retrieval.")
            return []
        if embeddings is None or text_chunks is None:
            st.warning("Required data files are missing. Skipping DSM-5 criteria retrieval.")
            return []

        try:
            print("Generating patient data embedding...")
            patient_embedding = voyage_client.embed([patient_data], model="voyage-3").embeddings[0]
            
            print("Finding initial relevant criteria...")
            similarities = cosine_similarity([patient_embedding], embeddings)[0]
            top_indices = similarities.argsort()[-initial_top_k:][::-1]
            
            initial_criteria = []
            for idx in top_indices:
                criteria_text = text_chunks[idx]
                initial_criteria.append({
                    'criteria': criteria_text,
                    'score': float(similarities[idx]),
                    'text': criteria_text
                })
            
            print("Reranking criteria...")
            rerank_input = [item['text'] for item in initial_criteria]
            reranked_results = voyage_client.rerank(patient_data, rerank_input, model="rerank-2", top_k=final_top_k)
            
            relevant_criteria = []
            for result in reranked_results.results:
                criteria = next(item for item in initial_criteria if item['text'] == result.document)
                relevant_criteria.append({
                    'criteria': criteria['criteria'],
                    'score': result.relevance_score
                })
            
            criteria_text = f"Found {len(relevant_criteria)} relevant criteria:\n"
            for i, criteria in enumerate(relevant_criteria, 1):
                # Split the criteria text into lines and take the first line
                first_line = criteria['criteria'].split('\n')[0].strip()
                criteria_text += f"{i}. {first_line} (Score: {criteria['score']:.4f})\n"
            print(criteria_text)
            self.add_to_document('Relevant DSM-5 Criteria', criteria_text, level=2)
            
            return relevant_criteria
        except Exception as e:
            st.error(f"Error in get_relevant_dsm5_criteria: {str(e)}")
            return []

    def generate_haiku_prompt(self, patient_data: str, relevant_criteria: List[Dict]):
        print("Generating Haiku prompt using Opus...")
        message = f"""
        You are tasked with creating a prompt for an AI model to analyze patient data and suggest possible diagnoses based on DSM-5 criteria. Your prompt should guide the AI to perform a thorough diagnostic analysis. Here's the information you need to incorporate:

Create a prompt that instructs the AI model to:

1. Carefully review the patient data provided above.
2. Compare the patient's symptoms with the DSM-5 criteria provided.
3. Identify which specific criteria are met, citing relevant patient information for each.
4. Suggest possible diagnoses based on the criteria met.
5. Provide a detailed explanation of the reasoning behind each suggested diagnosis.
6. Highlight any areas where more information might be needed for a more accurate diagnosis.

Now, generate a prompt for the AI model. 
Your prompt should emphasize the importance of a systematic approach and clear, evidence-based reasoning. Instruct the AI to structure its response as follows:

1. <criteria_analysis>
   List each relevant DSM-5 criterion and whether it is met, with supporting evidence from the patient data.
</criteria_analysis>

2. <possible_diagnoses>
   List each possible diagnosis, ordered from most to least likely.
</possible_diagnoses>

3. <diagnostic_reasoning>
   For each possible diagnosis, provide a detailed explanation of why it is being suggested, referencing specific criteria met and patient symptoms.
</diagnostic_reasoning>

4. <information_gaps>
   Identify any areas where additional information would be helpful for a more accurate diagnosis, explaining why this information is needed.
</information_gaps>
        
        Patient Data:
        {patient_data}

        Relevant DSM-5 Criteria:
        {json.dumps(relevant_criteria, indent=2)}

The prompt should be written as if addressing the AI model directly, instructing it to be thorough in its analysis, considering all provided information. 
        """

        response = anthropic_client.messages.create(
            model=self.opus_model,
            max_tokens=1000,
            messages=[
                {"role": "user", "content": message}
            ]
        )
        haiku_prompt = response.content[0].text
        self.add_to_document('Generated Haiku Prompt (Diagnosis)', haiku_prompt, level=2)
        return haiku_prompt

    def diagnose(self, patient_data: str):
        relevant_criteria = self.get_relevant_dsm5_criteria(patient_data)
        haiku_prompt = self.generate_haiku_prompt(patient_data, relevant_criteria)
        
        print("\nGenerated prompt for Haiku (Diagnosis):")
        print(haiku_prompt)
        
        print("\nGenerating diagnosis using Haiku...")
        response = anthropic_client.messages.create(
            model=self.haiku_model,
            max_tokens=1000,
            messages=[
                {"role": "user", "content": f"Patient Data:\n{patient_data}\n\n{haiku_prompt}"}
            ]
        )
        diagnosis = response.content[0].text
        full_diagnosis = f"Patient Data:\n{patient_data}\n\nDiagnosis:\n{diagnosis}"
        self.add_to_document('Diagnosis', full_diagnosis, level=2)
        return full_diagnosis

    def generate_treatment_prompt(self, patient_data: str, patient_history: str, diagnosis: str):
        print("Generating treatment prompt using Opus...")
        message = f"""
        You are tasked with creating a specific prompt for an AI model to suggest a treatment plan based on patient information. This prompt will be used to guide the AI in providing comprehensive and tailored treatment recommendations.

Now, generate a prompt for the AI model to suggest a treatment plan. Your prompt should instruct the AI to:

1. Consider appropriate therapy options
2. Suggest potential medications, if applicable
3. Recommend lifestyle changes that might be beneficial
4. Provide reasoning for each suggestion
5. Highlight any potential risks or side effects

Ensure that your prompt is specific, detailed, and tailored to the patient's data, history, and diagnosis. The prompt should guide the AI to provide a comprehensive treatment plan that takes into account l factors.
        {diagnosis}

        Patient History:
        {patient_history}

The prompt should be written as if addressing the AI model directly, instructing it on how to formulate the treatment plan based on the given information.
        """

        response = anthropic_client.messages.create(
            model=self.opus_model,
            max_tokens=1000,
            messages=[
                {"role": "user", "content": message}
            ]
        )
        treatment_prompt = response.content[0].text
        self.add_to_document('Generated Haiku Prompt (Treatment)', treatment_prompt, level=2)
        return treatment_prompt

    def suggest_treatment(self, patient_data: str, patient_history: str, diagnosis: str):
        treatment_prompt = self.generate_treatment_prompt(patient_data, patient_history, diagnosis)
        
        print("\nGenerated prompt for Haiku (Treatment):")
        print(treatment_prompt)
        
        print("\nGenerating treatment plan using Haiku...")
        response = anthropic_client.messages.create(
            model=self.haiku_model,
            max_tokens=1000,
            messages=[
                {"role": "user", "content": treatment_prompt}
            ]
        )
        treatment_plan = response.content[0].text
        self.add_to_document('Treatment Plan', treatment_plan, level=2)
        return treatment_plan

    def evaluate_final_analysis(self, final_analysis: str) -> tuple:
        print("\nEvaluating final analysis...")
        evaluation_prompt = f"""
        Evaluate the following final analysis of a patient's diagnosis and treatment plan. Determine if further iteration is needed based on the following criteria:
        1. Are there any areas identified as needing more detailed assessment?
        2. Are there suggestions for additional evaluations or interventions?
        3. Is there a clear indication that the current plan may not be fully appropriate or effective?

        If any of these criteria are met, recommend continuing the process with a focus on the identified areas.
        If no further iteration is needed, provide suggestions for enhancement.

        Final Analysis:
        {final_analysis}

        Provide your evaluation as a JSON object with the following structure:
        {{
            "continue_iteration": boolean,
            "reason": "string explaining the reason or focus areas if applicable, or suggestions for enhancement if no further iteration is needed"
        }}
        Ensure that your response contains only the JSON object, with no additional text before or after.
        """

        response = anthropic_client.messages.create(
            model=self.opus_model,
            max_tokens=500,
            messages=[
                {"role": "user", "content": evaluation_prompt}
            ]
        )
        evaluation = response.content[0].text
        print(f"Evaluation result: {evaluation}")
        
        # Parse the evaluation result
        try:
            evaluation_dict = json.loads(evaluation)
            return (evaluation_dict['continue_iteration'], evaluation_dict['reason'])
        except json.JSONDecodeError:
            print("Error parsing evaluation result. Defaulting to no further iteration.")
            return (False, "Error in evaluation: " + evaluation)

    def process_patient(self, patient_data: str, patient_history: str, max_iterations: int = 3):
        # We'll keep these in memory but not add them to the document
        self.patient_data = patient_data
        self.patient_history = patient_history
        
        iteration = 1
        full_diagnosis = self.diagnose(patient_data)
        treatment_plan = self.suggest_treatment(patient_data, patient_history, full_diagnosis)

        while iteration <= max_iterations:
            print(f"\nIteration {iteration}")

            if iteration > 1:
                # For subsequent iterations, focus on the areas identified for improvement
                improvement_prompt = f"""
                Based on the previous analysis, focus on the following areas for improvement:
                {reason}

                Provide updated recommendations for the diagnosis and treatment plan, addressing these specific areas.
                Previous diagnosis:
                {full_diagnosis}

                Previous treatment plan:
                {treatment_plan}
                """
                
                print("\nGenerating focused improvements using Opus...")
                improvement_response = anthropic_client.messages.create(
                    model=self.opus_model,
                    max_tokens=1000,
                    messages=[
                        {"role": "user", "content": improvement_prompt}
                    ]
                )
                improvements = improvement_response.content[0].text
                
                full_diagnosis += f"\n\nIteration {iteration} improvements:\n{improvements}"
                treatment_plan += f"\n\nIteration {iteration} improvements:\n{improvements}"

            print("\nGenerating final response using Opus...")
            final_prompt = f"""
            Analyze the following diagnosis and treatment plan, and provide a comprehensive summary for a healthcare provider. Your analysis should:
            1. Highlight key points from both the diagnosis and treatment plan
            2. Identify potential challenges in implementing the treatment plan
            3. Suggest any areas that may require further attention or assessment
            4. Provide an overall evaluation of the coherence and appropriateness of the diagnosis and treatment plan
            5. Include a comprehensive treatment plan and potential diagnoses analysis

            Patient Data:
            {self.patient_data}

            Patient History:
            {self.patient_history}

            Diagnosis:
            {full_diagnosis}

            Treatment Plan:
            {treatment_plan}

            Provide a detailed analysis and summary of the diagnosis and treatment plan, ensuring all relevant information is included.
            """

            final_response = anthropic_client.messages.create(
                model=self.opus_model,
                max_tokens=1500,
                messages=[
                    {"role": "user", "content": final_prompt}
                ]
            )

            final_analysis = final_response.content[0].text
            
            # We'll only keep the last iteration's final analysis
            if iteration == max_iterations:
                self.document = Document()  # Reset the document
                self.document.add_heading('Diagnostic and Treatment Report', 0)
                self.add_to_document('Final Analysis', final_analysis, level=1)

            continue_iteration, reason = self.evaluate_final_analysis(final_analysis)
            print(f"\nProcess completed iteration {iteration}.")
            print("\nAI Recommendation:")
            print(f"Continue iteration: {'Yes' if continue_iteration else 'No'}")
            print("Reason or Suggestions for Enhancement:")
            print(reason)

            user_input = input("\nDo you want to continue the process for another iteration? (yes/no): ").lower()
            if user_input != 'yes':
                break

            iteration += 1

        # Save the document
        self.document.save('Diagnostic_and_Treatment_Report.docx')
        print("\nReport saved as 'Diagnostic_and_Treatment_Report.docx'")
        
        return {
            "diagnosis": full_diagnosis,
            "treatment_plan": treatment_plan,
            "final_analysis": final_analysis
        }

def process_patient_with_progress(agent, patient_data, patient_history, max_iterations=3):
    iteration = 1
    status_area = st.empty()
    
    status_area.write("Starting diagnostic process...")
    full_diagnosis = agent.diagnose(patient_data)
    status_area.write("Initial diagnosis complete. Generating treatment plan...")
    treatment_plan = agent.suggest_treatment(patient_data, patient_history, full_diagnosis)

    while iteration <= max_iterations:
        update_progress(iteration, max_iterations)
        status_area.write(f"Processing iteration {iteration} of {max_iterations}")

        if iteration > 1:
            status_area.write("Generating focused improvements...")
            improvement_prompt = f"""
            Based on the previous analysis, focus on the following areas for improvement:
            {reason}

            Provide updated recommendations for the diagnosis and treatment plan, addressing these specific areas.
            Previous diagnosis:
            {full_diagnosis}

            Previous treatment plan:
            {treatment_plan}
            """
            
            improvement_response = anthropic_client.messages.create(
                model=agent.opus_model,
                max_tokens=1000,
                messages=[
                    {"role": "user", "content": improvement_prompt}
                ]
            )
            improvements = improvement_response.content[0].text
            
            full_diagnosis += f"\n\nIteration {iteration} improvements:\n{improvements}"
            treatment_plan += f"\n\nIteration {iteration} improvements:\n{improvements}"

        status_area.write("Generating final analysis...")
        final_prompt = f"""
        Analyze the following diagnosis and treatment plan, and provide a comprehensive summary for a healthcare provider. Your analysis should:
        1. Highlight key points from both the diagnosis and treatment plan
        2. Identify potential challenges in implementing the treatment plan
        3. Suggest any areas that may require further attention or assessment
        4. Provide an overall evaluation of the coherence and appropriateness of the diagnosis and treatment plan
        5. Include a comprehensive treatment plan and potential diagnoses analysis

        Patient Data:
        {patient_data}

        Patient History:
        {patient_history}

        Diagnosis:
        {full_diagnosis}

        Treatment Plan:
        {treatment_plan}

        Provide a detailed analysis and summary of the diagnosis and treatment plan, ensuring all relevant information is included.
        """

        final_response = anthropic_client.messages.create(
            model=agent.opus_model,
            max_tokens=1500,
            messages=[
                {"role": "user", "content": final_prompt}
            ]
        )

        final_analysis = final_response.content[0].text
        
        if iteration == max_iterations:
            agent.document = Document()
            agent.document.add_heading('Diagnostic and Treatment Report', 0)
            agent.add_to_document('Final Analysis', final_analysis, level=1)

        status_area.write("Evaluating final analysis...")
        continue_iteration, reason = agent.evaluate_final_analysis(final_analysis)
        
        if not continue_iteration:
            status_area.write("Process complete. Preparing final report...")
            break

        status_area.write("Preparing for next iteration...")
        iteration += 1

    agent.document.save('Diagnostic_and_Treatment_Report.docx')
    status_area.write("Report generated and saved.")
    
    return {
        "diagnosis": full_diagnosis,
        "treatment_plan": treatment_plan,
        "final_analysis": final_analysis
    }

# Streamlit app
st.title("Diagnostic and Treatment Assistant")

# Input fields
patient_data = st.text_area("Patient Data", height=200)
patient_history = st.text_area("Patient History", height=200)

if st.button("Process Patient"):
    if patient_data and patient_history:
        agent = DiagnosticTreatmentAgent()
        
        # Create a placeholder for the progress bar
        progress_bar = st.progress(0)
        
        # Create a placeholder for the status message
        status_message = st.empty()

        # Function to update progress
        def update_progress(iteration, max_iterations):
            progress = (iteration / max_iterations) * 100
            progress_bar.progress(int(progress))
            status_message.text(f"Processing iteration {iteration} of {max_iterations}")

        # Process the patient data
        with st.spinner('Processing patient data...'):
            result = process_patient_with_progress(agent, patient_data, patient_history)

        # Display the results
        st.success("Processing complete!")
        st.subheader("Final Analysis")
        st.write(result["final_analysis"])

        # Provide a download link for the Word document
        doc_buffer = io.BytesIO()
        agent.document.save(doc_buffer)
        doc_buffer.seek(0)
        st.download_button(
            label="Download Full Report",
            data=doc_buffer,
            file_name="Diagnostic_and_Treatment_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("Please provide both patient data and history.")